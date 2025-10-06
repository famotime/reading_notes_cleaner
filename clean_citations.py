from pathlib import Path
import docx
import re
from docx.oxml.ns import qn
from docx.oxml import OxmlElement # 虽然在提供的代码片段中未直接使用，但保留以防其他地方需要
from enum import Enum
import tempfile
import os
import sys

# 剪贴板支持
try:
    import pyperclip
    PYPERCLIP_AVAILABLE = True
except ImportError:
    PYPERCLIP_AVAILABLE = False
    print("警告: pyperclip 未安装，剪贴板处理功能受限")

# 导入 .doc 文件处理相关库
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False
    print("警告: python-docx2txt 未安装，.doc 文件将使用备用方法处理")

# Windows 平台的 COM 接口支持
if sys.platform == "win32":
    try:
        import win32com.client
        WIN32COM_AVAILABLE = True
    except ImportError:
        WIN32COM_AVAILABLE = False
        print("警告: pywin32 未安装，.doc 文件转换功能受限")
else:
    WIN32COM_AVAILABLE = False


class FileFormat(Enum):
    """文件格式枚举"""
    WORD = "word"
    MARKDOWN = "markdown"


def convert_doc_to_docx(doc_path):
    """
    将 .doc 文件转换为临时的 .docx 文件

    Args:
        doc_path: .doc 文件路径

    Returns:
        临时 .docx 文件路径，如果转换失败则返回 None
    """
    if not WIN32COM_AVAILABLE:
        return None

    try:
        # 创建临时文件
        temp_dir = tempfile.gettempdir()
        temp_docx_path = os.path.join(temp_dir, f"temp_{Path(doc_path).stem}.docx")

        # 使用 COM 接口打开 Word 应用程序
        word_app = win32com.client.Dispatch("Word.Application")
        word_app.Visible = False

        try:
            # 打开 .doc 文件
            doc = word_app.Documents.Open(str(Path(doc_path).absolute()))

            # 另存为 .docx 格式
            doc.SaveAs2(temp_docx_path, FileFormat=16)  # 16 = wdFormatXMLDocument (.docx)
            doc.Close()

            return temp_docx_path

        finally:
            # 确保关闭 Word 应用程序
            word_app.Quit()

    except Exception as e:
        print(f"使用 COM 接口转换 .doc 文件失败: {e}")
        return None


def read_doc_file_as_text(doc_path):
    """
    从 .doc 文件中提取纯文本内容（备用方法）

    Args:
        doc_path: .doc 文件路径

    Returns:
        提取的文本内容，如果失败则返回 None
    """
    if not DOCX2TXT_AVAILABLE:
        return None

    try:
        # 使用 docx2txt 提取文本（注意：这个库实际上也支持 .doc 文件）
        text = docx2txt.process(str(doc_path))
        return text
    except Exception as e:
        print(f"使用 docx2txt 提取 .doc 文件文本失败: {e}")
        return None


def clean_markdown_text(text):
    """
    清理Markdown文本中的脚注和引用

    Args:
        text: Markdown文本内容

    Returns:
        清理后的文本
    """
    # 移除方括号引用 [1], [2-4] 等
    text = re.sub(r'\[\d+(?:-\d+)?\]', '', text)
    # 移除 cite 风格的引用标记，如 [cite: xxx]、[cite_start]
    text = re.sub(r'\[cite(?:_[a-z]+)?(?::[^\]]*)?\]', '', text, flags=re.IGNORECASE)
    # 移除脚注数字（如 1. １。2、3, 4．等，含前后空格）
    # 注意：这个正则表达式可能会移除段落开头的合法编号，如果需要保留，需要更精确的匹配
    text = re.sub(r'\s*[\d０-９]+[\.。．、,，]?\s*', '', text) # 这个正则比较宽泛
    # 去除多余空格
    text = re.sub(r'[ \t]+', ' ', text)
    # 移除标点前多余空格
    text = re.sub(r'\s+([,.;:，。；：！？!?])', r'\1', text)
    return text.strip()


def remove_footnotes_citations(doc):
    """
    删除文档中的所有脚注和引用，仅在非标题且非列表段落清理序号
    """
    for paragraph in doc.paragraphs:
        is_heading = paragraph.style.name.startswith('Heading') or paragraph.style.name.startswith('标题')
        is_list = 'List' in paragraph.style.name or '项目符号' in paragraph.style.name or '编号' in paragraph.style.name
        for run in paragraph.runs:
            if hasattr(run._element, 'xpath'):
                footnote_refs = run._element.xpath('.//w:footnoteReference')
                for ref in footnote_refs:
                    ref.getparent().remove(ref)
            text = run.text
            # 只在非标题且非列表段落清理序号
            if not is_heading and not is_list:
                # 移除方括号引用，如[1], [2-4]等
                text = re.sub(r'\[\d+(?:-\d+)?\]', '', text)
                # 移除 cite 风格的引用标记，如 [cite: xxx]、[cite_start]
                text = re.sub(r'\[cite(?:_[a-z]+)?(?::[^\]]*)?\]', '', text, flags=re.IGNORECASE)
                # 移除数字序号（包括中英文数字和标点），如"1. "、"１。"、"2、"等，要求后面是空格或结尾
                text = re.sub(r'(?<![\d０-９\.])[\d０-９]+[\.。、,，]?(?=\s|$)', '', text)
                # 将多个连续的空格或制表符替换为单个空格
                text = re.sub(r'[ \t]+', ' ', text)
            # 移除标点前多余空格
            text = re.sub(r'\s+([,.;:，。；：！？!?])', r'\1', text)
            run.text = text.strip() # 注意：这里的strip()可能会导致只包含空格的run变为空
    # 表格同理
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # 表格单元格内的段落通常没有独立的标题或列表样式，
                    # 但保留此检查以保持与上方逻辑的一致性（尽管其影响可能较小）
                    is_heading_cell_para = paragraph.style.name.startswith('Heading') or paragraph.style.name.startswith('标题')
                    is_list_cell_para = 'List' in paragraph.style.name or '项目符号' in paragraph.style.name or '编号' in paragraph.style.name
                    for run in paragraph.runs:
                        if hasattr(run._element, 'xpath'):
                            footnote_refs = run._element.xpath('.//w:footnoteReference')
                            for ref in footnote_refs:
                                ref.getparent().remove(ref)
                        text = run.text
                        if not is_heading_cell_para and not is_list_cell_para:
                            text = re.sub(r'\[\d+(?:-\d+)?\]', '', text)
                            text = re.sub(r'(?<![\d０-９\.])[\d０-９]+[\.。、,，]?(?=\s|$)', '', text)
                        text = re.sub(r'[ \t]+', ' ', text)
                        run.text = text.strip() # 同样，strip()可能清空run
    return doc


def get_paragraph_numbering(paragraph):
    """
    获取Word段落的自动编号（如2.4、1.等），如无则返回空字符串。
    """
    p = paragraph._p
    numPr = p.find(qn('w:numPr'))
    if numPr is None:
        return ''

    # 尝试从段落文本中提取编号作为一种回退或简化方法
    # python-docx本身不直接提供编号的文本值
    # 这里的实现是基于文本模式匹配，可能不适用于所有复杂的编号格式
    text = paragraph.text.strip()
    # 匹配 "1." "1．" "1、" "1," "１." 等简单列表编号
    m = re.match(r'^([\d０-９]+(?:[\.。．、,，](?=\s|$))?)', text)
    if m:
        # 提取编号部分，并尝试移除编号后的文本部分以便后续内容处理
        num_text = m.group(1)
        # 确保编号后确实是列表项内容，而不是普通段落恰好以数字开头
        # （此简化版本主要依赖后续的 is_list 判断）
        return num_text.strip()

    # 兼容嵌套编号如 "2.4" "2.5." 等，后面可能跟空格或直接是内容
    m2 = re.match(r'^([\d０-９]+(?:\.[\d０-９]+)*[\.。．、,，]?(?=\s|$))', text)
    if m2:
        num_text = m2.group(1)
        return num_text.strip()

    return ''


def extract_text_with_formatting(doc):
    """
    提取文档内容，保留基本格式、加粗、超链接，转换为Markdown文本，支持有序/无序列表、自动编号和表格。
    """
    markdown_text = []
    current_position = 0  # 用于跟踪当前处理位置

    # 辅助函数：从run中提取文本，支持加粗、斜体
    def get_run_markdown(run):
        run_text = run.text
        if run.bold:
            run_text = f"**{run_text}**"
        if run.italic:
            run_text = f"*{run_text}*"
        return run_text

    # 辅助函数：提取段落中的超链接和普通文本，返回完整Markdown字符串
    def get_paragraph_markdown_with_links(para):
        md_parts = []
        rels = para.part.rels
        # 遍历底层xml，识别超链接
        for child in para._p:
            if child.tag.endswith('hyperlink'):
                # 处理超链接
                rId = child.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                url = rels[rId].target_ref if rId in rels else None
                link_text = ''
                for r in child.findall('.//w:r', namespaces=child.nsmap):
                    # 处理加粗、斜体
                    run_obj = None
                    for run in para.runs:
                        if hasattr(run, '_element') and run._element is r:
                            run_obj = run
                            break
                    if run_obj:
                        link_text += get_run_markdown(run_obj)
                    else:
                        # fallback
                        link_text += ''.join(t.text or '' for t in r.findall('.//w:t', namespaces=r.nsmap))
                if url:
                    md_parts.append(f'[{link_text}]({url})')
                else:
                    md_parts.append(link_text)
            elif child.tag.endswith('r'):
                # 普通run
                run_obj = None
                for run in para.runs:
                    if hasattr(run, '_element') and run._element is child:
                        run_obj = run
                        break
                if run_obj:
                    md_parts.append(get_run_markdown(run_obj))
                else:
                    # fallback
                    texts = child.findall('.//w:t', namespaces=child.nsmap)
                    md_parts.append(''.join(t.text or '' for t in texts))
            # 其他类型节点忽略
        return ''.join(md_parts)

    # 辅助函数：从单元格提取Markdown格式的文本（支持超链接）
    def get_cell_markdown(cell):
        cell_para_markdowns = []
        for para_in_cell in cell.paragraphs:
            para_text = get_paragraph_markdown_with_links(para_in_cell)
            if para_text.strip():
                cell_para_markdowns.append(para_text.strip())
        final_cell_text = "<br>".join(cell_para_markdowns)
        final_cell_text = final_cell_text.replace("|", "\\|")
        return final_cell_text.strip()

    # 辅助函数：判断是否为无序列表
    def is_unordered_list(para):
        style_name = para.style.name.lower()
        if 'bullet' in style_name or '项目符号' in style_name:
            return True
        text = para.text.lstrip()
        if text.startswith(('•', '·', '●', '-', '*', '○', '■', '□', '◆', '◇')):
            return True
        if hasattr(para._p, 'pPr') and para._p.pPr is not None:
            numPr = para._p.pPr.find(qn('w:numPr'))
            if numPr is not None:
                numId = numPr.find(qn('w:numId'))
                if numId is not None:
                    try:
                        num_id = int(numId.get(qn('w:val')))
                        if num_id < 100:
                            return True
                    except (ValueError, TypeError):
                        pass
        return False

    # 辅助函数：处理无序列表项，保留格式
    def process_unordered_list_item(para):
        # 获取完整的格式化文本
        full_text = get_paragraph_markdown_with_links(para)

        # 找到第一个非项目符号的字符位置
        text = para.text.lstrip()
        bullet_match = re.match(r'^[•·●\-*○■□◆◇]\s*', text)
        if bullet_match:
            # 计算项目符号的长度（包括可能的空格）
            bullet_length = len(bullet_match.group(0))
            # 保留项目符号后的所有格式
            return full_text[bullet_length:].lstrip()
        return full_text

    # 获取所有段落和表格的位置信息
    elements = []
    for element in doc.element.body:
        if element.tag.endswith('p'):
            elements.append(('p', element))
        elif element.tag.endswith('tbl'):
            elements.append(('tbl', element))

    for element_type, element in elements:
        if element_type == 'p':
            para = doc.paragraphs[current_position]
            current_position += 1

            style_name = para.style.name
            is_heading1 = style_name.startswith('Heading 1') or style_name == '标题 1'
            is_heading2 = style_name.startswith('Heading 2') or style_name == '标题 2'
            is_heading3 = style_name.startswith('Heading 3') or style_name == '标题 3'
            is_list_style = 'List' in style_name or '项目符号' in style_name or '编号' in style_name

            numbering_text = get_paragraph_numbering(para)
            is_actual_list_item = para._p.pPr is not None and para._p.pPr.numPr is not None

            line_prefix = ""
            content_to_append = ""

            if is_heading1 or is_heading2 or is_heading3:
                formatted_text = get_paragraph_markdown_with_links(para)
                line_prefix = '#' * (1 if is_heading1 else 2 if is_heading2 else 3) + ' '
                if numbering_text and formatted_text.startswith(numbering_text):
                    content_to_append = formatted_text[len(numbering_text):].lstrip()
                    line_prefix += f"{numbering_text.strip()} "
                else:
                    content_to_append = formatted_text

            elif is_actual_list_item or is_list_style:
                if is_unordered_list(para):
                    content_to_append = process_unordered_list_item(para)
                    line_prefix = "- "
                else:
                    # 有序列表处理
                    formatted_text = get_paragraph_markdown_with_links(para)
                    cleaned_content = formatted_text
                    if numbering_text and formatted_text.startswith(numbering_text):
                        cleaned_content = formatted_text[len(numbering_text):].lstrip()
                        line_prefix = f"{numbering_text.strip()} "
                    else:
                        num_match = re.match(r'^([\d０-９]+[\.。．、,，])\s*', formatted_text)
                        if num_match:
                            num = num_match.group(1).replace('。', '.').replace('．', '.').replace('、', '.').replace('，', '.').replace(',', '.')
                            cleaned_content = formatted_text[len(num_match.group(0)):].lstrip()
                            line_prefix = f"{num.strip()} "
                        else:
                            line_prefix = "1. "
                    content_to_append = cleaned_content

            else:
                # 普通段落
                content_to_append = get_paragraph_markdown_with_links(para)

            if not content_to_append.strip() and not markdown_text:
                continue
            if not content_to_append.strip() and markdown_text and markdown_text[-1] == "":
                continue
            if not content_to_append.strip():
                markdown_text.append("")
                continue

            markdown_text.append(f"{line_prefix}{content_to_append}")

        elif element_type == 'tbl':
            table = None
            for t in doc.tables:
                if t._element is element:
                    table = t
                    break

            if not table or not table.rows:
                continue

            table_md_lines = []
            header_cells_md = []

            if table.rows:
                first_row_cells = table.rows[0].cells
                for cell in first_row_cells:
                    header_text = get_cell_markdown(cell)
                    header_cells_md.append(header_text)

                table_md_lines.append("| " + " | ".join(header_cells_md) + " |")
                table_md_lines.append("| " + " | ".join(["---"] * len(header_cells_md) if header_cells_md else ["---"]) + " |")

            start_row_index = 1 if header_cells_md else 0
            for row_idx in range(start_row_index, len(table.rows)):
                row = table.rows[row_idx]
                row_cells_md = []
                for cell in row.cells:
                    cell_text = get_cell_markdown(cell)
                    row_cells_md.append(cell_text)

                if header_cells_md:
                    if len(row_cells_md) < len(header_cells_md):
                        row_cells_md.extend([""] * (len(header_cells_md) - len(row_cells_md)))
                    elif len(row_cells_md) > len(header_cells_md):
                        row_cells_md = row_cells_md[:len(header_cells_md)]

                table_md_lines.append("| " + " | ".join(row_cells_md) + " |")

            if table_md_lines:
                markdown_text.append("\n" + "\n".join(table_md_lines))

    return "\n\n".join(markdown_text).strip()


def convert_word_to_markdown(word_path, output_dir=None, overwrite=False):
    """
    将Word文档（.doc 或 .docx）转换为markdown，清理脚注和引用

    Args:
        word_path: Word文档路径（支持 .doc 和 .docx 格式）
        output_dir: 输出目录，默认为Word文档所在目录
        overwrite: 是否覆盖已存在的文件，默认为False

    Returns:
        output_file: 输出文件路径，如果文件已存在且不覆盖则返回None
    """
    word_path = Path(word_path)
    output_dir = Path(output_dir) if output_dir else word_path.parent
    output_dir.mkdir(exist_ok=True, parents=True)

    output_file = output_dir / f"{word_path.stem}.md"

    if output_file.exists() and not overwrite:
        print(f"跳过已存在的文件 (不覆盖): {output_file}")
        return None

    # 判断文件格式并处理
    file_extension = word_path.suffix.lower()
    temp_docx_path = None

    try:
        if file_extension == '.docx':
            # 直接处理 .docx 文件
            doc = docx.Document(word_path)
            doc = remove_footnotes_citations(doc)  # 先清理
            markdown_content = extract_text_with_formatting(doc)  # 再提取

        elif file_extension == '.doc':
            # 处理 .doc 文件
            print(f"处理 .doc 文件: {word_path}")

            # 方法1: 尝试转换为 .docx 然后处理（保留格式）
            temp_docx_path = convert_doc_to_docx(word_path)
            if temp_docx_path and os.path.exists(temp_docx_path):
                print(f"成功转换为临时 .docx 文件，保留完整格式")
                doc = docx.Document(temp_docx_path)
                doc = remove_footnotes_citations(doc)
                markdown_content = extract_text_with_formatting(doc)
            else:
                # 方法2: 备用方案，提取纯文本
                print(f"使用备用方案提取纯文本")
                text_content = read_doc_file_as_text(word_path)
                if text_content:
                    # 对纯文本应用基本的引用清理
                    markdown_content = clean_markdown_text(text_content)
                else:
                    raise Exception("无法读取 .doc 文件内容")
        else:
            raise Exception(f"不支持的文件格式: {file_extension}")

        # 写入输出文件
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(markdown_content)

        return output_file

    finally:
        # 清理临时文件
        if temp_docx_path and os.path.exists(temp_docx_path):
            try:
                os.remove(temp_docx_path)
            except Exception as e:
                print(f"清理临时文件失败: {e}")


# 保持向后兼容性的别名函数
def convert_docx_to_markdown(docx_path, output_dir=None, overwrite=False):
    """
    向后兼容的函数，调用新的统一处理函数
    """
    return convert_word_to_markdown(docx_path, output_dir, overwrite)


def convert_markdown_to_cleaned(md_path, output_dir=None, overwrite=False):
    """
    清理Markdown文件中的脚注和引用

    Args:
        md_path: Markdown文件路径
        output_dir: 输出目录，默认为Markdown文件所在目录
        overwrite: 是否覆盖已存在的文件，默认为False

    Returns:
        output_file: 输出文件路径，如果文件已存在且不覆盖则返回None
    """
    md_path = Path(md_path)
    output_dir = Path(output_dir) if output_dir else md_path.parent
    output_dir.mkdir(exist_ok=True, parents=True)

    output_file = output_dir / f"{md_path.stem}_cleaned.md"

    if output_file.exists() and not overwrite:
        print(f"跳过已存在的文件 (不覆盖): {output_file}")
        return None

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    cleaned_content = clean_markdown_text(content) # 使用基础的Markdown清理

    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(cleaned_content)

    return output_file



def process_clipboard_markdown(output_path=None, copy_back=True, overwrite=False):
    """
    从剪贴板读取Markdown文本并进行清理

    Args:
        output_path: 可选输出文件路径
        copy_back: 是否将清理结果写回剪贴板
        overwrite: 写入文件时是否覆盖已存在的文件

    Returns:
        cleaned_text: 清理后的Markdown文本
    """
    if not PYPERCLIP_AVAILABLE:
        print("错误: pyperclip 未安装，无法处理剪贴板内容。")
        return None

    try:
        clipboard_text = pyperclip.paste()
    except pyperclip.PyperclipException as e:
        print(f"读取剪贴板失败: {e}")
        return None

    if not clipboard_text:
        print("剪贴板为空或不包含文本，跳过处理。")
        return None

    cleaned_text = clean_markdown_text(clipboard_text)

    if output_path:
        output_path = Path(output_path)
        if output_path.exists() and not overwrite:
            print(f"跳过写入 (不覆盖): {output_path}")
        else:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(cleaned_text, encoding='utf-8')
            print(f"已将清理后的文本写入: {output_path}")

    if copy_back:
        try:
            pyperclip.copy(cleaned_text)
            print("已将清理结果写回剪贴板。")
        except pyperclip.PyperclipException as e:
            print(f"写回剪贴板失败: {e}")

    print("清理后的Markdown内容:\n")
    print(cleaned_text)

    return cleaned_text



def process_directory(input_dir, output_dir=None, overwrite=False, file_format=FileFormat.WORD):
    """
    处理目录下所有文件，包括子目录

    Args:
        input_dir: 输入目录
        output_dir: 输出目录，如果不指定，则保存在原文件同一目录下
        overwrite: 是否覆盖已存在的文件，默认为False
        file_format: 文件格式，支持Word和Markdown，默认为Word
    """
    input_dir = Path(input_dir)
    processed_output_dir = None
    if output_dir: # 只有当用户明确指定output_dir时才使用
        processed_output_dir = Path(output_dir)
        processed_output_dir.mkdir(exist_ok=True, parents=True)

    if file_format == FileFormat.WORD:
        # 搜索 .doc 和 .docx 文件，排除临时文件
        docx_files = [f for f in input_dir.rglob("*.docx") if not f.name.startswith("~$")]
        doc_files = [f for f in input_dir.rglob("*.doc") if not f.name.startswith("~$")]
        files = docx_files + doc_files
        convert_func = convert_word_to_markdown
    elif file_format == FileFormat.MARKDOWN:
        files = list(input_dir.rglob("*.md"))
        convert_func = convert_markdown_to_cleaned
    else:
        print(f"不支持的文件格式: {file_format}")
        return

    if not files:
        print(f"目录 {input_dir} 及其子目录中没有找到 {file_format.value} 文件。")
        return

    total = len(files)
    success = 0
    skipped = 0
    failed = 0

    for file_path in files:
        try:
            current_output_dir = file_path.parent # 默认输出到原文件目录
            if processed_output_dir: # 如果指定了全局输出目录
                # 计算相对路径，并在全局输出目录下创建相同的子目录结构
                relative_path = file_path.relative_to(input_dir)
                current_output_dir = processed_output_dir / relative_path.parent
                current_output_dir.mkdir(exist_ok=True, parents=True)

            output_file = convert_func(file_path, current_output_dir, overwrite)
            if output_file:
                print(f"已处理: {file_path} -> {output_file}")
                success += 1
            else:
                skipped += 1
        except Exception as e:
            print(f"处理 {file_path} 时出错: {e}")
            failed += 1

    print("\n处理统计结果：")
    print(f"  总文件数: {total}")
    print(f"  成功: {success}")
    print(f"  跳过: {skipped}")
    print(f"  失败: {failed}")


if __name__ == "__main__":
    # --- 配置区 ---
    # INPUT_DIR_PATH = Path("input")  
    INPUT_DIR_PATH = None  # 设置为 None 可跳过目录批量处理

    # 可选：指定统一的输出目录。如果为 None，则输出到各原文件所在目录。
    # OUTPUT_DIR_PATH = Path("./DeepResearch Reports/Markdown_Output")
    OUTPUT_DIR_PATH = None

    OVERWRITE_EXISTING = False  # 是否覆盖已存在的输出文件

    # 选择处理的文件类型：FileFormat.WORD 或 FileFormat.MARKDOWN
    PROCESSING_FORMAT = FileFormat.WORD

    # 剪贴板处理开关及设置
    PROCESS_CLIPBOARD_MARKDOWN = True
    CLIPBOARD_OUTPUT_PATH = None  # 例如 Path("clipboard_cleaned.md")
    CLIPBOARD_OVERWRITE = False
    CLIPBOARD_COPY_BACK = True
    # --- 配置区结束 ---

    if PROCESS_CLIPBOARD_MARKDOWN:
        print("开始处理剪贴板中的Markdown文本。")
        process_clipboard_markdown(
            output_path=CLIPBOARD_OUTPUT_PATH,
            copy_back=CLIPBOARD_COPY_BACK,
            overwrite=CLIPBOARD_OVERWRITE
        )
        print("剪贴板处理完成。")

    if INPUT_DIR_PATH is None:
        print("未配置输入目录，跳过批量文件处理。")
    else:
        if not INPUT_DIR_PATH.exists() or not INPUT_DIR_PATH.is_dir():
            print(f"错误：输入目录 '{INPUT_DIR_PATH}' 不存在或不是一个目录。")
        else:
            print(f"开始处理目录: {INPUT_DIR_PATH}")
            print(f"处理文件类型: {PROCESSING_FORMAT.value}")
            if OUTPUT_DIR_PATH:
                print(f"输出到指定目录: {OUTPUT_DIR_PATH}")
            else:
                print("输出到原文件各自所在目录 (或其子目录，如_cleaned)。")
            print(f"是否覆盖已存在文件: {OVERWRITE_EXISTING}")

            process_directory(
                INPUT_DIR_PATH,
                output_dir=OUTPUT_DIR_PATH,
                overwrite=OVERWRITE_EXISTING,
                file_format=PROCESSING_FORMAT
            )
            print("处理完成。")

