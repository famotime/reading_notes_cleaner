"""
合并markdown/docx文件内容，有两种模式：
1. 将目录下的所有markdown/docx文件内容，按照标题重新组合，相同标题下的内容拼接在一起，合并为一个markdown文件
2. 将指定目录下以及各子目录的markdown/docx文件内容，分别合并为一个markdown文件，原文件名作为二级标题，目录名_combined作为新文件名
"""
from pathlib import Path
import re
from docx import Document


def combine_markdown_files_by_title(directory):
    """将指定目录下的所有markdown/docx文件内容，按照标题重新组合，相同标题下的内容拼接在一起，合并为一个markdown文件"""
    content_by_header = {}

    # Regex to match headers (# to ######)
    header_pattern = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)

    dir_path = Path(directory)

    # 处理markdown文件
    for file_path in dir_path.glob('*.md'):
        # 排除包含"_combined"的文件
        if '_combined' in file_path.stem:
            continue

        content = file_path.read_text(encoding='utf-8')

        sections = header_pattern.split(content)

        for i in range(1, len(sections), 3):
            level = sections[i]
            header = sections[i+1]
            text = sections[i+2]

            full_header = f"{level} {header}"

            if full_header not in content_by_header:
                content_by_header[full_header] = []

            content_by_header[full_header].append(text.strip())

    # 处理docx文件
    for file_path in dir_path.glob('*.docx'):
        # 排除包含"_combined"的文件
        if '_combined' in file_path.stem:
            continue

        doc = Document(file_path)
        content = []
        current_header = None
        current_text = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 检查是否是标题
            header_match = header_pattern.match(text)
            if header_match:
                # 保存之前的内容
                if current_header and current_text:
                    if current_header not in content_by_header:
                        content_by_header[current_header] = []
                    content_by_header[current_header].append('\n\n'.join(current_text))

                # 开始新的标题
                current_header = text
                current_text = []
            else:
                if current_header:
                    current_text.append(text)
                else:
                    # 没有标题的内容放在默认标题下
                    default_header = "# 未分类内容"
                    if default_header not in content_by_header:
                        content_by_header[default_header] = []
                    content_by_header[default_header].append(text)

        # 保存最后一个标题的内容
        if current_header and current_text:
            if current_header not in content_by_header:
                content_by_header[current_header] = []
            content_by_header[current_header].append('\n\n'.join(current_text))

    combined_content = []
    for header, contents in content_by_header.items():
        combined_content.append(header)
        combined_content.append("\n\n\n\n".join(contents))
        combined_content.append("\n\n\n\n")

    output_path = dir_path / f"{dir_path.stem}_combined.md"
    output_path.write_text("\n".join(combined_content), encoding='utf-8')


def combine_markdown_files(directory):
    """将指定目录下以及各子目录的markdown和docx文件内容，分别合并为一个markdown文件，原文件名作为二级标题，目录名_combined作为新文件名"""
    dir_path = Path(directory)

    def natural_sort_key(path):
        """自然排序的key函数，用于处理文件名中的数字"""
        import re
        # 将文件名中的数字转换为整数进行比较
        convert = lambda text: int(text) if text.isdigit() else text.lower()
        return [convert(c) for c in re.split('([0-9]+)', path.stem)]

    # 遍历目录及子目录
    for folder in dir_path.glob('**/'):
        if folder == dir_path:
            continue

        # 获取当前目录下所有markdown和docx文件，排除包含"_combined"的文件，并按自然顺序排序
        md_files = sorted([f for f in folder.glob('*.md') if '_combined' not in f.stem], key=natural_sort_key)
        docx_files = sorted([f for f in folder.glob('*.docx') if '_combined' not in f.stem], key=natural_sort_key)
        if not (md_files or docx_files):
            continue

        # 创建images子目录
        images_dir = folder / 'images'
        images_dir.mkdir(exist_ok=True)

        # 合并内容
        combined_content = []

        # 处理markdown文件
        for md_file in md_files:
            # 添加文件名作为二级标题
            combined_content.append(f"## {md_file.stem}")
            combined_content.append("\n")

            # 添加文件内容
            content = md_file.read_text(encoding='utf-8')
            combined_content.append(content.strip())
            combined_content.append('\n\n')

        # 处理docx文件
        for docx_file in docx_files:
            # 添加文件名作为二级标题
            combined_content.append(f"## {docx_file.stem}")
            combined_content.append("\n")

            # 读取docx内容并转换为markdown
            doc = Document(docx_file)
            content = []

            # 处理段落和图片
            for element in doc.element.body.iter():
                # 处理图片
                if element.tag.endswith('drawing'):
                    for img in element.iter():
                        if img.tag.endswith('blip'):
                            # 获取图片关系ID
                            rId = img.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if rId:
                                # 获取图片二进制数据
                                image_part = doc.part.related_parts[rId]
                                image_bytes = image_part.blob

                                # 生成唯一的图片文件名
                                image_filename = f"{docx_file.stem}_{rId}{Path(image_part.partname).suffix}"
                                image_path = images_dir / image_filename

                                # 保存图片
                                image_path.write_bytes(image_bytes)

                                # 添加markdown图片链接
                                relative_path = f"images/{image_filename}"
                                content.append(f"![{image_filename}]({relative_path})\n")

                # 处理文本段落
                elif element.tag.endswith('p'):
                    text = ''.join(t.text for t in element.iter() if t.text)
                    if text.strip():
                        content.append(text.strip())

            combined_content.append('\n\n'.join(content))
            combined_content.append('\n\n')

        # 保存合并后的文件
        output_path = folder / f"{folder.name}_combined.md"
        output_path.write_text('\n'.join(combined_content), encoding='utf-8')


if __name__ == "__main__":
    directory = r"D:\小汤汁茶馆知识星球哈"
    # combine_markdown_files_by_title(directory)
    combine_markdown_files(directory)
